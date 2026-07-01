from docx import Document

from app.analyzers.docx import analyze_docx


def test_analyze_docx_basic(tmp_path):
	docx_path = tmp_path / "test.docx"

	doc = Document()
	doc.core_properties.author = "Test Author"
	doc.core_properties.title = "Test Document"
	doc.add_paragraph("Hello World")
	doc.save(str(docx_path))

	result = analyze_docx(str(docx_path))

	assert result["file_name"] == "test.docx"
	assert result["type"] == "docx"
	assert result["author"] == "Test Author"
	assert result["paragraphs"] == 1


def test_analyze_docx_empty(tmp_path):
	docx_path = tmp_path / "test.docx"

	doc = Document()
	doc.save(str(docx_path))

	result = analyze_docx(str(docx_path))

	assert result["type"] == "docx"
	assert result["paragraphs"] == 0


def test_analyze_docx_no_file():
	import pytest

	with pytest.raises(FileNotFoundError):
		analyze_docx("nonexistent.docx")


def test_analyze_docx_fields(tmp_path):
	docx_path = tmp_path / "fields.docx"

	doc = Document()
	doc.add_paragraph("Hello World")
	doc.add_table(rows=2, cols=3)
	doc.save(str(docx_path))

	result = analyze_docx(str(docx_path))

	assert "has_comments" in result
	assert "styles_count" in result
	assert "images_count" in result
	assert "tables_count" in result
	assert "sections" in result
	assert result["tables_count"] == 1
