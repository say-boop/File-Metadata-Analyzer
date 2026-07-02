from docx import Document
from pathlib import Path


def analyze_docx(file_path: str) -> dict:
	path_obj = Path(file_path)

	result = {
		"file_name": path_obj.name,
		"file_size": path_obj.stat().st_size,
		"path": path_obj,
		"type": "docx",
		"author": "N/A",
		"created": "N/A",
		"modified": "N/A",
		"last_modified_by": "N/A",
		"revision": "N/A",
		"paragraphs": "N/A",
		"has_comments": False,
		"styles_count": 0,
		"images_count": 0,
		"tables_count": 0,
		"sections": 0,
	}
	
	try:
		doc = Document(file_path)
		props = doc.core_properties

		result["author"] = props.author or "N/A"
		result["created"] = str(props.created) if props.created else "N/A"
		result["modified"] = str(props.modified) if props.modified else "N/A"
		result["last_modified_by"] = props.last_modified_by or "N/A"
		result["revision"] = props.revision or "N/A"
		result["paragraphs"] = len(doc.paragraphs)
		result["has_comments"] = len(doc.comments) > 0 if doc.comments else False
		result["styles_count"] = len(doc.styles)
		result["tables_count"] = len(doc.tables)
		result["sections"] = len(doc.sections)

		try:
			images = 0

			for para in doc.paragraphs:
				for run in para.runs:
					if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
						images += 1

					if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pic"):
						images += 1

			result["images_count"] = images
		except Exception:
			pass

		return result
	except Exception:
		return result
